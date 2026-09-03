# -*- coding: utf-8 -*-
"""把 使用说明书.md 渲染成排版受控的 使用说明书.docx。

排版规则：A4 / 边距 2.5cm；正文微软雅黑 11pt、1.5 倍行距、
段后 6pt；标题三级全部黑色加粗（不用 Word 内置 Heading 样式，避免默认蓝绿色）；
列表用普通段落 + 缩进 + 手动编号；行内 `code` 与目录树代码块用 Consolas 等宽、
东亚字体微软雅黑；粗体由 md 里的 ** 控制（已大幅削减）。

用法：python gen_manual_docx.py [md路径] [docx输出路径]
默认：使用说明书.md -> 使用说明书.docx（当前目录）
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

FONT = "微软雅黑"
MONO = "Consolas"
BLACK = (0, 0, 0)

BODY_SIZE = 11
TITLE_SIZE = 20
H2_SIZE = 14
H3_SIZE = 12
CODE_SIZE = 10.5

DEFAULT_MD = "使用说明书.md"
DEFAULT_OUT = "使用说明书.docx"


def set_run_font(run, ascii_font, east_font, size, bold, color):
    """给 run 设置西文/中文/复杂脚本字体、字号、加粗、颜色（全手动，不依赖模板）。"""
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:ascii"), ascii_font)
    rFonts.set(qn("w:hAnsi"), ascii_font)
    rFonts.set(qn("w:eastAsia"), east_font)
    rFonts.set(qn("w:cs"), ascii_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor(*color)


def _para_fmt(p, before, after, line, indent=None, align=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line
    if indent is not None:
        pf.left_indent = Cm(indent)
    if align is not None:
        pf.alignment = align


def add_inline(p, text, size=BODY_SIZE, base_bold=False):
    """解析 **粗体** 与 `行内代码`，逐段生成 run。"""
    segs = text.split("**")
    bold = base_bold
    for seg in segs:
        if not seg:
            bold = not bold
            continue
        parts = seg.split("`")
        for j, part in enumerate(parts):
            if not part:
                continue
            run = p.add_run(part)
            if j % 2 == 1:
                set_run_font(run, MONO, FONT, size, bold, BLACK)
            else:
                set_run_font(run, FONT, FONT, size, bold, BLACK)
        bold = not bold


def set_shading(p, fill="F2F2F2"):
    pPr = p._element.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


def setup_page(doc):
    sec = doc.sections[0]
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5)
    sec.right_margin = Cm(2.5)
    # 文档默认字体兜底（Normal 样式）
    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(BODY_SIZE)
    normal.font.color.rgb = RGBColor(*BLACK)


def render_title(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, 0, 12, 1.3, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, FONT, FONT, TITLE_SIZE, True, BLACK)


def render_heading(doc, text, size, before, after):
    p = doc.add_paragraph()
    _para_fmt(p, before, after, 1.2)
    add_inline(p, text, size=size, base_bold=True)


def render_body(doc, text):
    p = doc.add_paragraph()
    _para_fmt(p, 0, 6, 1.5)
    add_inline(p, text)


def render_list_item(doc, text, prefix):
    p = doc.add_paragraph()
    _para_fmt(p, 0, 3, 1.5, indent=0.74)
    add_inline(p, prefix + text)


def render_code_block(doc, code_lines):
    for line in code_lines:
        p = doc.add_paragraph()
        _para_fmt(p, 0, 0, 1.0, indent=0.4)
        set_shading(p)
        run = p.add_run(line)
        set_run_font(run, MONO, FONT, CODE_SIZE, False, BLACK)


def render(md_path, out_path):
    doc = Document()
    setup_page(doc)
    with open(md_path, "r", encoding="utf-8") as f:
        lines = f.read().splitlines()

    i, n = 0, len(lines)
    while i < n:
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue
        if stripped.startswith("```"):
            i += 1
            code_lines = []
            while i < n and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i].rstrip())
                i += 1
            i += 1  # 跳过收尾 ```
            render_code_block(doc, code_lines)
            continue
        if stripped.startswith("### "):
            render_heading(doc, stripped[4:].strip(), H3_SIZE, 10, 6)
            i += 1
            continue
        if stripped.startswith("## "):
            render_heading(doc, stripped[3:].strip(), H2_SIZE, 18, 8)
            i += 1
            continue
        if stripped.startswith("# "):
            render_title(doc, stripped[2:].strip())
            i += 1
            continue
        m = re.match(r"^[-*•]\s+(.*)$", stripped)
        if m:
            render_list_item(doc, m.group(1), "•  ")
            i += 1
            continue
        m = re.match(r"^(\d+)[.、）)]\s*(.*)$", stripped)
        if m:
            render_list_item(doc, m.group(2), f"{m.group(1)}. ")
            i += 1
            continue
        render_body(doc, stripped)
        i += 1

    doc.save(out_path)
    return out_path


def main(argv=None):
    args = argv if argv is not None else sys.argv[1:]
    md = args[0] if len(args) > 0 else DEFAULT_MD
    out = args[1] if len(args) > 1 else DEFAULT_OUT
    path = render(md, out)
    print("OK ->", path)


if __name__ == "__main__":
    main()
