# -*- coding: utf-8 -*-
"""会议纪要 Markdown → docx（python-docx）+ 同名 Markdown 原样落盘。

排版只做三档：标题（#/##/###）/ 正文 / 列表（- 无序、1. 有序），不做复杂渲染。
成对导出（save_minutes_pair）：会议纪要_YYYYMMDD_HHMM.docx / .md 同名成对，
存输出目录；docx 页脚写生成时间，md 为 generate_minutes 返回的原始 Markdown。
save_minutes_docx 保留（只存 docx，兼容旧调用）。
"""
import os
import re
import time

from docx import Document
from docx.shared import Pt

_BULLET_RE = re.compile(r"^\s*[-*•]\s+(.+)$")
_NUMBER_RE = re.compile(r"^\s*\d+[.、）)]\s*(.+)$")


def _plain(text):
    """剥掉行内残留符号（反引号/下划线强调），供标题用。"""
    return text.replace("`", "").replace("__", "").replace("**", "")


def _add_text_paragraph(doc, text, style=None):
    """正文/列表段落：支持 **粗体** 行内标记，其余符号剥掉。"""
    p = doc.add_paragraph(style=style)
    bold = False
    for seg in text.split("**"):
        seg = seg.replace("`", "").replace("__", "")
        if not seg:
            continue
        run = p.add_run(seg)
        run.bold = bold
        bold = not bold
    if not p.runs:
        p.add_run("")
    return p


def _render_markdown(doc, md):
    """Markdown 行 → docx 段落（标题/正文/列表三档）。"""
    for raw in md.splitlines():
        line = raw.rstrip()
        if not line.strip():
            continue
        if line.startswith("### "):
            doc.add_heading(_plain(line[4:].strip()), level=3)
        elif line.startswith("## "):
            doc.add_heading(_plain(line[3:].strip()), level=2)
        elif line.startswith("# "):
            doc.add_heading(_plain(line[2:].strip()), level=1)
        else:
            m = _BULLET_RE.match(line)
            if m:
                _add_text_paragraph(doc, m.group(1), style="List Bullet")
                continue
            m = _NUMBER_RE.match(line)
            if m:
                _add_text_paragraph(doc, m.group(1), style="List Number")
                continue
            _add_text_paragraph(doc, line.strip())


def _save_docx_at(markdown, path):
    """把纪要 Markdown 渲染成 docx 存到指定路径（标题/正文/列表三档，页脚写时间）。"""
    full = time.strftime("%Y年%m月%d日 %H:%M")
    doc = Document()
    doc.add_heading("会议纪要", level=0)
    meta = doc.add_paragraph()
    run = meta.add_run(f"生成时间：{full}")
    run.font.size = Pt(10)
    _render_markdown(doc, markdown)
    footer = doc.sections[0].footer
    footer.paragraphs[0].text = f"生成时间：{full}"
    doc.save(path)


def _next_base_name(out_dir, stamp):
    """下一个可用的文件名基（无扩展名）：docx 与 md 都不存在的第一个序号。

    基名 会议纪要_<stamp>，冲突则 会议纪要_<stamp>_2 / _3…（两种扩展名任一
    占用都跳过，保证成对文件永远同基名）。
    """
    base = os.path.join(out_dir, f"会议纪要_{stamp}")
    n = 2
    while os.path.exists(f"{base}.docx") or os.path.exists(f"{base}.md"):
        base = os.path.join(out_dir, f"会议纪要_{stamp}_{n}")
        n += 1
    return base


def save_minutes_pair(markdown, out_dir):
    """纪要成对导出：docx + md 同名成对放 out_dir。返回 (docx_path, md_path)。

    - 文件名基 = 会议纪要_YYYYMMDD_HHMM（同名冲突时两文件同基名加 _2/_3…）；
    - md 内容 = generate_minutes 返回的原始 Markdown 原样写（utf-8、不加 BOM，
      文件头不加任何多余东西）；
    - docx 渲染复用 _save_docx_at（与 save_minutes_docx 同一套排版）。
    """
    os.makedirs(out_dir, exist_ok=True)
    base = _next_base_name(out_dir, time.strftime("%Y%m%d_%H%M"))
    docx_path = base + ".docx"
    md_path = base + ".md"
    with open(md_path, "w", encoding="utf-8", newline="\n") as f:
        f.write(markdown)
    _save_docx_at(markdown, docx_path)
    return docx_path, md_path


def save_minutes_docx(markdown, out_dir):
    """把纪要 Markdown 存成 docx，返回文件绝对路径。同名自动加序号。"""
    stamp = time.strftime("%Y%m%d_%H%M")
    os.makedirs(out_dir, exist_ok=True)
    path = os.path.join(out_dir, f"会议纪要_{stamp}.docx")
    n = 2
    while os.path.exists(path):
        path = os.path.join(out_dir, f"会议纪要_{stamp}_{n}.docx")
        n += 1
    _save_docx_at(markdown, path)
    return path
